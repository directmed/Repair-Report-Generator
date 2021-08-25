from tkinter import *
from tkinter import messagebox
from tkinter.filedialog import askdirectory
from selenium.common.exceptions import NoSuchElementException, InvalidSessionIdException

# SETTING UP tkinter
root = Tk()
root.iconbitmap('DM_icon.ico')
root.title('Repair Report Generator')


# Options Page
def open_options():
    options_window = Toplevel()
    options_window.title('Options')
    options_window.iconbitmap('DM_icon.ico')
    user_options_entries = []
    options_entries_text = [StringVar(), StringVar(), StringVar()]

    def on_closing():
        message_response = messagebox.askyesnocancel(title=None, message='Do you want to save?')
        if message_response is True:  # save options then close window
            try:
                print(type(int(options_entries_text[0].get())))
                # Startup wait time - wait time for user to enter their information
                if int(options_entries_text[1].get()) <= 44:
                    # print(int(options_entries_text[1].get()))
                    options_entries_text[1].set('45')

                # Order wait time - Time the program will wait for information to load between each internet page
                if int(options_entries_text[2].get()) <= 0:
                    # print(int(options_entries_text[2].get()))
                    options_entries_text[2].set('1')

                user_options[0] = options_entries_text[0].get()
                user_options[1] = options_entries_text[1].get()
                user_options[2] = options_entries_text[2].get()
                save_options()
                reset_status_box()
                update_status_box('Options have been saved.')
                options_window.destroy()

            except ValueError:
                # print('Recognized ValueError')
                messagebox.showerror('Value Error', 'Please enter Integers only.')

        elif message_response is False:
            # close window without saving
            options_window.destroy()
        else:
            # do nothing
            pass

    for index, text in enumerate(user_options):
        options_entries_text[index].set(text)

    # INFORMATION LABELS
    option_labels = [Label(options_window, text='RON link index:'),
                     Label(options_window, text='Startup Wait Time:'),
                     Label(options_window, text='Report Wait Time:')]

    # SETTING LABELS' POSITIONS AND DIMENSIONS
    for option_index, option_label in enumerate(option_labels):
        option_label.grid(row=option_index, column=0, padx=80)
        user_options_entries.append(Entry(options_window, width=40, textvariable=options_entries_text[option_index]))
        user_options_entries[option_index].grid(row=option_index, column=1)

    options_window.grab_set()  # places this window in front of root window.
    options_window.protocol("WM_DELETE_WINDOW", on_closing)


# Options Page
def open_xpath():
    xpath_window = Toplevel()
    xpath_window.title('Xpath')
    xpath_window.iconbitmap('DM_icon.ico')
    xpath_options_entries = []
    xpath_entries_text = [StringVar(), StringVar(), StringVar(), StringVar(), StringVar(),
                          StringVar(), StringVar(), StringVar(), StringVar(), StringVar()]

    def on_closing():
        message_response = messagebox.askyesnocancel(title=None, message='Do you want to save?')
        if message_response is True:  # save options then close window
            try:
                xpath_options[0] = xpath_entries_text[0].get()
                xpath_options[1] = xpath_entries_text[1].get()
                xpath_options[2] = xpath_entries_text[2].get()
                xpath_options[3] = xpath_entries_text[3].get()
                xpath_options[4] = xpath_entries_text[4].get()
                xpath_options[5] = xpath_entries_text[5].get()
                xpath_options[6] = xpath_entries_text[6].get()
                xpath_options[7] = xpath_entries_text[7].get()
                xpath_options[8] = xpath_entries_text[8].get()
                xpath_options[9] = xpath_entries_text[9].get()

                save_xpaths()
                reset_status_box()
                update_status_box('Xpaths have been saved.')
                xpath_window.destroy()

            except ValueError:
                # print('Recognized ValueError')
                messagebox.showerror('Value Error', 'Please enter Integers only.')

        elif message_response is False:
            # close window without saving
            xpath_window.destroy()
        else:
            # do nothing
            pass

    for index, text in enumerate(xpath_options):
        xpath_entries_text[index].set(text)

    # INFORMATION LABELS
    xpath_labels = [Label(xpath_window, text='Coil Make:'),
                    Label(xpath_window, text='Part Number:'),
                    Label(xpath_window, text='Serial Number:'),
                    Label(xpath_window, text='Customer Complaint:'),
                    Label(xpath_window, text='Initial Failure:'),
                    Label(xpath_window, text='Repair Notes:'),
                    Label(xpath_window, text='Sales Order:'),
                    Label(xpath_window, text='Customer Contact:'),
                    Label(xpath_window, text='Product Description:'),
                    Label(xpath_window, text='Customer Name:')]

    # SETTING LABELS' POSITIONS AND DIMENSIONS
    for xpath_index, xpath_label in enumerate(xpath_labels):
        xpath_label.grid(row=xpath_index, column=0, padx=80)
        xpath_options_entries.append(Entry(xpath_window, width=120, textvariable=xpath_entries_text[xpath_index]))
        xpath_options_entries[xpath_index].grid(row=xpath_index, column=1)

    xpath_window.grab_set()  # places this window in front of root window.
    xpath_window.protocol("WM_DELETE_WINDOW", on_closing)


# pass user_info
def directory_button():
    this_directory = askdirectory(parent=root)
    if this_directory != '':
        user_entries[3].delete(0, 'end')
        user_entries[3].insert(0, this_directory)


def reset_status_box():
    status_box.config(state='normal')
    status_box.delete("1.0", "end")
    status_box.config(state='disabled')


def update_repair_button_text():
    this_state = chkValue.get()
    if this_state is True:
        report_button.config(text='Run Startup')
        dir_button.config(state='disabled')
        # load_button.config(state='disabled')
        # save_button.config(state='disabled')
        user_entries[2].config(state='disabled')
        user_entries[3].config(state='disabled')
        user_entries[4].config(state='disabled')
        reset_status_box()
        update_status_box('Run Startup:\n\n* Will delete all of the saved cookies if any.\n\nMake sure your '
                          'information is correct.\n\n* User must go through Salesforce.com sign in process. This may '
                          'include verifying your account via text message or through some other means. You will have '
                          + user_options[1] + ' seconds to complete the verification process.\n\n\nYou can extend '
                                              'the Waiting time in the "Options Menu". Minimum 45 seconds.')

    elif this_state is False:
        report_button.config(text='Generate Report')
        dir_button.config(state='normal')
        # load_button.config(state='normal')
        # save_button.config(state='normal')
        user_entries[2].config(state='normal')
        user_entries[3].config(state='normal')
        user_entries[4].config(state='normal')
        reset_status_box()
        update_status_box('Generate Report:\nMake sure the selected directory is correct. You can change this at any '
                          'time.\n\nMake sure that all of the information is correct.\n\nPress *Load* to load saved '
                          'data.\n\nPress *Save* to save currently displayed data.')


def update_status_box(message):
    status_box.config(state='normal')
    status_box.insert(END, message)
    status_box.config(state='disabled')


def checkbox_state():
    these_messages = ['User name is empty. Please enter your Salesforce.com account name.\n',
                      'Password is empty. Please enter your Salesforce.com password.\n',
                      'The Engineer name is empty. Please enter your name.\n',
                      'Select a directory.\n']
    empty_var = [StringVar(), StringVar(), StringVar(), StringVar()]
    empty_var[3].set('C:')
    this_state = chkValue.get()
    if this_state is True:
        for msg_index, this_message in enumerate(these_messages):
            if msg_index == 2:
                break

            if entries_text[msg_index].get() == empty_var[msg_index].get():
                reset_status_box()
                update_status_box(this_message)
                return

        start_up()
    else:
        for msg_index, this_message in enumerate(these_messages):
            if entries_text[msg_index].get() == empty_var[msg_index].get():
                reset_status_box()
                update_status_box(this_message)
                return

        generate_report()


# pass user_info
def load_info():
    this_flag = False
    try:
        with open('user_info.txt') as f:
            f.seek(0)
            first_char = f.read(1)
            if not first_char:
                f.seek(0)
                this_flag = True
                reset_status_box()
                update_status_box('Save file is empty.\n\nPlease go through the setup process by checking the box\n\n'
                                  'OR\n\nEnter your information and press the *Save* button.')
            else:
                f.seek(0)
                data = f.read()
                this_info = data.split(",")
                info_index = 0
                for data, entry in zip(this_info, user_entries):
                    # entry = Entry(root)
                    entry.delete(0, 'end')
                    entry.insert(0, data)
                    # entry.grid(row=index, column=1)
                    info_index = info_index + 1

                reset_status_box()
                update_status_box('Data has been loaded.')
                this_flag = False
    except FileNotFoundError:
        this_flag = True

    if this_flag is True:
        save_info()
        return 0


# pass user_options
def load_options():
    this_flag = False
    try:
        with open('user_options.txt') as f:
            f.seek(0)
            first_char = f.read(1)
            if not first_char:
                this_flag = True

            else:
                f.seek(0)
                data = f.read()
                this_info = data.split(",")
                options_index = 0
                for data in this_info:
                    # entry = Entry(root)
                    user_options[options_index] = data
                    # print(data)
                    # print(user_options[options_index])
                    # entry.grid(row=index, column=1)
                    options_index = options_index + 1
                    this_flag = False
    except FileNotFoundError:
        this_flag = True

    if this_flag is True:
        user_options[0] = '6'
        user_options[1] = '120'
        user_options[2] = '3'
        save_options()
        return 0


# pass user_options
def load_xpaths():
    this_flag = False
    try:
        with open('xpath_options.txt') as f:
            f.seek(0)
            first_char = f.read(1)
            if not first_char:
                this_flag = True

            else:
                f.seek(0)
                data = f.read()
                this_info = data.split(",")
                options_index = 0
                for data in this_info:
                    # entry = Entry(root)
                    xpath_options[options_index] = data
                    # print(data)
                    # print(user_options[options_index])
                    # entry.grid(row=index, column=1)
                    options_index = options_index + 1
                    this_flag = False
    except FileNotFoundError:
        this_flag = True

    if this_flag is True:
        # Coil Make
        xpath_options[0] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[5]/div[1]/div/form/div[2]/table/tbody/tr[2]/th'
        # Part Number
        xpath_options[1] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[5]/div[1]/div/form/div[2]/table/tbody/tr[2]/td[3]'
        # Serial Number
        xpath_options[2] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[5]/div[1]/div/form/div[2]/table/tbody/tr[2]/td[4]'
        # Customer Complaint
        xpath_options[3] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[8]/table/tbody/tr[3]/td[2]/div/table/tbody/tr/td/div'
        # Initial Failure
        xpath_options[4] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[8]/table/tbody/tr[4]/td[2]/div'
        # Repair Notes
        xpath_options[5] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[8]/table/tbody/tr[5]/td[2]/div'
        # Sales Order
        xpath_options[6] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[2]/table/tbody/tr[2]/td[4]/div/a'
        # Customer Contact
        xpath_options[7] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[5]/table/tbody/tr[4]/td[2]/div'
        # Product Description
        xpath_options[8] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[5]/div[1]/div/form/div[2]/table/tbody/tr'
        # Customer Name
        xpath_options[9] = '/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[5]/table/tbody/tr[22]/td[2]/div'

        save_xpaths()
        return 0


# Pass user_info
def save_info():
    with open('user_info.txt', 'w') as f:
        # user_name, user_password, engineer_name, directory
        info_str = entries_text[0].get() + ',' + entries_text[1].get() + ',' + entries_text[2].get() + ',' + \
                   entries_text[3].get()
        f.write(info_str)

    reset_status_box()
    update_status_box('New data has been saved.')


# Pass user_info
def save_options():
    info_str = user_options[0] + ',' + user_options[1] + ',' + user_options[2]

    # print(info_str)
    # print(type(info_str))
    with open('user_options.txt', 'w') as f:
        f.write(info_str)

    reset_status_box()
    load_options()
    # update_status_box('New options have been saved.')


# Pass user_info
def save_xpaths():
    info_str = xpath_options[0] + ',' + xpath_options[1] + ',' + xpath_options[2] + ',' + xpath_options[3] + ',' + \
               xpath_options[4] + ',' + xpath_options[5] + ',' + xpath_options[6] + ',' + xpath_options[7] + ',' + \
               xpath_options[8] + ',' + xpath_options[9]
    # print(info_str)
    # print(type(info_str))
    with open('xpath_options.txt', 'w') as f:
        f.write(info_str)

    reset_status_box()
    load_xpaths()
    # update_status_box('New options have been saved.')


# Checks if popup comes up. Closes the popup if it does.
def check_for_popup(d):
    try:
        d.find_element_by_id('tryLexDialog').is_displayed()
        d.find_element_by_id('tryLexDialogX').click()
    except NoSuchElementException:
        return
    except InvalidSessionIdException:
        update_status_box('Window has been closed. Starting over.')


# --------------------------------------------------------------------------------
def start_up():
    from selenium import webdriver
    from selenium.common.exceptions import NoSuchElementException
    from selenium.webdriver.chrome.options import Options
    from selenium.common.exceptions import WebDriverException, NoSuchAttributeException
    import os

    # Set login info:  Set your user name and password for Salesforce.com
    user_name = entries_text[0].get()  # set the user_name form data in entry box
    user_password2 = entries_text[1].get()  # set the password form data in entry box
    # Delete all the messages in the status box
    status_box.config(state='normal')
    status_box.delete("1.0", "end")
    status_box.config(state='disabled')

    try:
        cookie_dir = os.path.abspath("chrome-data")
        cookie_dir_text = str(cookie_dir)
        chrome_options = Options()
        chrome_options.add_argument("--user-data-dir=" + cookie_dir_text)
        driver = webdriver.Chrome('chromedriver.exe', options=chrome_options)
        driver.implicitly_wait(int(user_options[1]))

        # Sign-in to SalesForce
        driver.get("https://directmedparts.my.salesforce.com/")
        elem = driver.find_element_by_id("username")
        elem.clear()
        elem.send_keys(user_name)
        elem = driver.find_element_by_id("password")
        elem.clear()
        elem.send_keys(user_password2)
        # elem = driver.find_element_by_id("Login").click()

        # Wait for user to verify information and click log in button
        driver.implicitly_wait(150)
        # check if there is a pop up
        check_for_popup(driver)

        if driver.find_element_by_class_name('searchBoxClearContainer').is_displayed():
            try:
                # Finding drop-down menu
                # print('Finding drop-down menu...')
                update_status_box('Finding drop-down menu...')
                elem = driver.find_element_by_xpath('/html/body/div[1]/div[1]/table/tbody/tr/td[3]/div/div[3]/div/div')
                elem.click()
                # finding Log out button
                update_status_box('clicking Logout button...')
                if driver.find_element_by_xpath('/html/body/div[1]/div[1]/table/tbody/tr/td[3]/div/div[3]/div/div/div['
                                                '2]/div[3]/a[4]').is_displayed():
                    driver.find_element_by_xpath(
                        '/html/body/div[1]/div[1]/table/tbody/tr/td[3]/div/div[3]/div/div/div[2]/div[3]/a[4]').click()
                # Reopening page to load new cookies. Signing in automatically to make sure it works.
                update_status_box('reopening page...')
                driver.get("https://directmedparts.my.salesforce.com/")
                update_status_box('entering user information....')
                if driver.find_element_by_xpath(
                        "/html/body/div[1]/div[1]/div/div/div[2]/div[3]/form/div[1]/div/div/div/a").is_displayed():
                    try:
                        if driver.find_element_by_id("password"):
                            try:
                                elem = driver.find_element_by_id("password")
                                elem.clear()
                                elem.send_keys(user_password2)
                                driver.find_element_by_id("Login").click()
                                update_status_box('cookies check 1: Info was saved...')

                            except NoSuchElementException:
                                update_status_box('*** Error:    Could not enter password....')

                    except NoSuchElementException:
                        if driver.find_element_by_id('username'):
                            try:
                                elem = driver.find_element_by_id("username")
                                elem.clear()
                                elem.send_keys(user_name)
                                elem = driver.find_element_by_id("password")
                                elem.clear()
                                elem.send_keys(user_password2)
                                driver.find_element_by_id("Login").click()
                                update_status_box('Cookies check 1: Info was NOT saved...')

                            except NoSuchElementException:
                                update_status_box('*** Could not log in again. Something is wrong....')
                            except NoSuchAttributeException:
                                update_status_box('No such attribute')

            except NoSuchElementException:
                update_status_box('*** Could not open HD 12 Channel Body Array (P Connector) page.')
        else:
            update_status_box('*** something went wrong with if statement.')
    except WebDriverException:
        reset_status_box()
        update_status_box('Chrome window was closed.')
        return

    save_info()
    update_status_box('Cookies have been saved.\nDirectMed user name and password have been saved.')


# ----------------------------------------------------------------------------------------------------
def generate_report():
    from selenium import webdriver
    from selenium.webdriver.common.keys import Keys
    from selenium.common.exceptions import NoSuchElementException
    from selenium.webdriver.chrome.options import Options
    from datetime import date
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import os
    import subprocess

    # Set login info:  Set your user name and password for Salesforce.com
    user_name = entries_text[0].get()  # set the user_name form data in entry box
    user_password2 = entries_text[1].get()  # set the password form data in entry box
    engineer = entries_text[2].get()  # set the engineer name from data in the entry box
    directory = entries_text[3].get()  # set the directory from data in the entry box
    ron_link_index_option = int(user_options[0])
    # Delete all the messages in the status box
    reset_status_box()
    today = date.today()  # current date

    # Prompt user for specific repair order number (ron), exit if not enough digits
    entries_text[4].set(int(entries_text[4].get()).__str__())
    if int(entries_text[4].get()) >= 100000 or int(entries_text[4].get()) <= 10000:
        update_status_box('Error:   Enter a valid Repair Order number.')
        return
    else:
        ron = entries_text[4].get()

    file_name = '/RepairOrder_' + ron.__str__()  # add repair order number to file name
    status_box.insert('end', 'Obtaining data from Salesforce.com')

    # Load cookies created in start_up.py: If user cannot log in, run start_up.py
    cookie_dir = os.path.abspath("chrome-data")
    cookie_dir_text = str(cookie_dir)
    chrome_options = Options()
    chrome_options.add_argument("--user-data-dir=" + cookie_dir_text)
    driver = webdriver.Chrome('chromedriver.exe', options=chrome_options)
    driver.implicitly_wait(user_options[2])
    # Log-in to Salesforce
    driver.get("https://directmedparts.my.salesforce.com/")

    # Check if user info is saved
    try:
        driver.find_element_by_xpath(
            "/html/body/div[1]/div[1]/div/div/div[2]/div[3]/form/div[1]/div/div/div/a").is_displayed()
        # print('user name was saved.')
    except NoSuchElementException:
        elem = driver.find_element_by_id("username")
        elem.clear()
        elem.send_keys(user_name)
        # print('user name NOT saved.')

    # Entering user name and password
    elem = driver.find_element_by_id("password")
    elem.clear()
    elem.send_keys(user_password2)
    driver.find_element_by_id("Login").click()
    check_for_popup(driver)

    # Check if Lightning Salesforce: switch to classic if so.

    # Search for the ron
    try:
        # Xpath - search bar
        elem = driver.find_element_by_xpath('/html/body/div[1]/div[1]/table/tbody/tr/td[2]/form/div/div/div[1]/input')
        elem.clear()
        elem.send_keys(ron)
        elem.send_keys(Keys.ENTER)
        # print('Searching for the repair order...')
    except NoSuchElementException:
        update_status_box('Error:   Search box not found.')
        driver.close()
        return
        # print('Search box not found.')

    # Find ron link in the results page. Click the link if found.
    try:
        parent = driver.find_elements_by_tag_name('th')
        for index, child in enumerate(parent):
            if index > ron_link_index_option:
                # print('index:  ' + index.__str__())
                # print('parent - 5:  ' + parent[index - ron_link_index_option].text)
                # print('child:  ' + child.text)
                if child.text == ron and parent[index - ron_link_index_option].text == 'Repair Order Number':
                    # print('found the right link!!')
                    grandchild = child.find_element_by_xpath("./child::*")
                    # print(grandchild.text)
                    grandchild.click()
                    break

    except NoSuchElementException:
        update_status_box('Error:   Could not find repair order in results.')
        driver.close()
        return
        # print('Could not find results list.')

    # Check for popups in RO page.
    check_for_popup(driver)

    # Date
    this_date = today.strftime("%m/%d/%Y")

    # Engineer Name
    repair_table[6] = 'Repair Engineer: ' + engineer
    # Coil Make
    elem = driver.find_element_by_xpath(xpath_options[0])
    # print('Coil Make:  ')
    # print(elem.text)
    info_table[1][1] = elem.text

    # Frequency
    if info_table[1][1].find('3') != -1:
        info_table[1][3] = '127MHz'
    elif info_table[1][1].find('1.5') != -1:
        info_table[1][3] = '64MHz'
    else:
        info_table[1][3] = '???'
        # print('Manually edit Frequency in document.')

    # Part Number
    elem = driver.find_element_by_xpath(xpath_options[1])
    # print('Part Number:  ')
    # print(elem.text)
    info_table[2][1] = elem.text

    # Serial Number
    elem = driver.find_element_by_xpath(xpath_options[2])
    # print('Serial Number:  ')
    # print(elem.text)
    info_table[2][3] = elem.text

    # Customer Complain
    elem = driver.find_element_by_xpath(xpath_options[3])
    # print('Customer complaint:  ')
    # print(elem.text)
    repair_table[1] = elem.text

    # Initial Failure
    elem = driver.find_element_by_xpath(xpath_options[4])
    # print('Initial failure:  ')
    # print(elem.text)
    repair_table[3] = elem.text

    # Repair notes
    elem = driver.find_element_by_xpath(xpath_options[5])
    # print('Repair notes:  ')
    # print(elem.text)
    repair_table[5] = elem.text

    '''
    # Go to product page to find the manufacturer
    driver.implicitly_wait(7)
    elem = driver.find_element_by_xpath('/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[5]/div[1]/div/form/div['
                                        '2]/table/tbody/tr[2]/th/a')
    elem.click()
    elem.find_element_by_xpath('/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[4]/div[2]/div[7]/table/tbody/tr[3]/td[2]/div')
    ro_info['manufacturer'] = elem.text
    driver.back()
    driver.implicitly_wait(3)
    '''
    # Go to sales order to find remaining information.
    try:
        elem = driver.find_element_by_xpath(xpath_options[6])
        # print('Sales Order:  ' + elem.text)
        # print('href:  ' + elem.get_property('href'))
        elem.click()  # click on SO link
    except NoSuchElementException:
        update_status_box('Error:   Could not find Sales Order.')
        driver.close()
        return

    # Contact
    elem = driver.find_element_by_xpath(xpath_options[7])
    # print('Contact:  ')
    # print(elem.text)
    info_table[0][3] = elem.text

    # Go to Product Description in SO to find remaining information.
    try:
        parent = driver.find_elements_by_xpath(xpath_options[8])
        for index, sibling in enumerate(parent):
            '''
            print('\n Parent index:  ' + index.__str__())
            print('sibling:  ')
            print(sibling.text)
            print('Sibling attributes:  ')
            print(sibling.get_property('attributes'))
            '''
            sibling_text = sibling.text
            if sibling_text.find(info_table[2][3]) != -1:
                child = driver.find_element_by_xpath('/html/body/div[1]/div[2]/table/tbody/tr/td[2]/div[5]/div['
                                                     '1]/div/form/div[2]/table/tbody/tr[' + (
                                                             index + 1).__str__() + ']/th/a')
                child.click()
                break
            else:
                continue

    except NoSuchElementException:
        update_status_box('Error:   No product description found.')
        driver.close()
        return
        # print('Could not find element.')

    try:
        # Customer Name
        elem = driver.find_element_by_xpath(xpath_options[9])
        # print('Customer Name:  ')
        # print(elem.text)
        info_table[0][1] = elem.text
    except NoSuchElementException:
        update_status_box('Error:   Could not find Account Name (Customer Name).')
        driver.close()
        return
        # print('Could not find Account Name (Customer Name).')

    driver.close()
    # print('Creating Report...')

    # CREATING DOCUMENT -------------------------------------------
    document = Document()
    # Adding Header
    section = document.sections[0]  # Header
    header = section.header
    # picture = document.add_picture('DM_Logo.jpg')
    paragraph = header.paragraphs[0]
    logo_run = paragraph.add_run()
    logo_run.add_picture('DM_Logo.jpg')
    text_run = paragraph.add_run()

    text_run.text = '\t\tDate: ' + this_date.__str__()
    paragraph.style = document.styles['Header']
    # Add Image

    # Adding Table 1:  Style = 'Plain Table 3'
    # Col_widths: col_1 = 1.25"= 1143000, col_2 = 3.07"= 2807208, col_3 = 1.00"= 914400, col_4 = 1.18"= 1078992
    # 32 chars max on 'coil make'
    col_width_table_1 = (1143000, 2807208, 914400, 1078992)
    col_width_table_2 = (1463040, 978408, 978408, 2523744)
    row_height_table_1 = (173736, 173736, 173736,)
    row_height_table_3 = (173736, 457200, 173736, 1133856, 173736, 1371600, 173736)
    table = document.add_table(3, 4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # row = 3, col = 4
    for row_index, info in enumerate(info_table):
        for col_index, detail in enumerate(info):
            table.cell(row_index, col_index).text = detail

    for col_index, width in enumerate(col_width_table_1):
        table.cell(0, col_index).width = width

    rows = table.rows
    for row_index, height in enumerate(row_height_table_1):
        rows[row_index].height = height

    document.add_paragraph(' ')
    # row = 9, col = 4
    table = document.add_table(9, 4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for row_index, info in enumerate(test_table):
        for col_index, detail in enumerate(info):
            table.cell(row_index, col_index).text = detail

    for col_index, width in enumerate(col_width_table_2):
        table.cell(0, col_index).width = width

    document.add_paragraph(' ')
    # row = 7, col = 1
    table = document.add_table(7, 1)
    table.style.hidden = False
    table.style = 'Medium List 1 Accent 1'

    for row_index, info in enumerate(repair_table):
        table.cell(row_index, 0).text = info

    rows = table.rows
    for row_index, height in enumerate(row_height_table_3):
        rows[row_index].height = height

    # Adding Footer
    section = document.sections[0]  # Header
    header = section.footer
    paragraph = header.paragraphs[0]
    paragraph.text = '\twww.DirectMedParts.com\r\t12525 Stowe Drive • Poway, CA • 92064'
    paragraph.style = document.styles['Footer']
    file_path = directory.__str__() + file_name.__str__() + '.docx'
    document.save(file_path)  # Save the document
    update_status_box(
        'Report complete. Check your directory.\n\nFile Path:\n' + file_path)
    subprocess.Popen(r'explorer /select,"' + file_path.replace('/', '\\') + '"')
    # ------------------------------------------------------------------------------------------


# ============================================================================================================
# DEFINING VARIABLES
# row = 3, col = 4
info_table = [['Customer Name:', ' ', 'Contact:', ' '],
              ['Coil Make:', ' ', 'Frequency:', ' '],
              ['Part Number:', ' ', 'Serial Number:', ' ']]
# row = 9, col = 4
test_table = [[' ', 'Incoming\rInspection', 'Outgoing\rInspection', 'Notes'],
              [' ', 'Pass/Fail', 'Pass/Fail', ' '],
              ['General Functionality', 'P', 'P', ' '],
              ['DC Check', 'P', 'P', ' '],
              ['Impedance Verification', ' ', ' ', ' '],
              ['S12 Verification', 'P', 'P', ' '],
              ['Uniformity', ' ', ' ', ' '],
              ['Decoupling Test', ' ', ' ', ' '],
              ['Other', 'P', 'P', ' ']]
# row = 7, col = 1
repair_table = ['Customer Complaint:',
                ' ',
                'Analysis of Findings:',
                ' ',
                'Repair Report:',
                ' ',
                'Repair Engineer: ']

# MENU BAR SETUP #
menu_bar = Menu(root)
root.config(menu=menu_bar)
# File menu
file_menu = Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label='File', menu=file_menu)
file_menu.add_command(label='Load Info', command=load_info)
file_menu.add_command(label='Save Info', command=save_info)
file_menu.add_separator()
file_menu.add_command(label='Exit', command=root.quit)
# Edit menu
edit_menu = Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label='Edit', menu=edit_menu)
edit_menu.add_command(label='Options', command=open_options)
edit_menu.add_command(label='Xpath', command=open_xpath)

# KTINKER VARIABLES #
# RON index, Startup Wait Time, Report Wait Time
user_options = [' ', ' ', ' ']
xpath_options = [' ', ' ', ' ', ' ', ' ', ' ', ' ', ' ', ' ', ' ']
user_entries = []
# Account Name, Account Password, Engineer Name, Directory, Repair Order Number
entries_text = [StringVar(), StringVar(), StringVar(), StringVar(), StringVar()]

# INFORMATION LABELS
labels = [Label(root, text='DirectMed User Name:   '),
          Label(root, text='DirectMed Password:    '),
          Label(root, text='Repair Engineer:       '),
          Label(root, text='Directory:             '),
          Label(root, text='Repair Order Number:   ')]

# SETTING LABELS' POSITIONS AND DIMENSIONS
for user_index, label in enumerate(labels):
    label.grid(row=user_index, column=0, padx=80)
    user_entries.append(Entry(root, width=40, textvariable=entries_text[user_index]))
    user_entries[user_index].grid(row=user_index, column=1)

# STATUS TEXTBOX LABEL STATUS AND POSITION
status_label = Label(root, text='Status')
status_box = Text(root)
status_box.config(state='disabled')
status_box.grid(row=5, column=1)

# MAKING PASSWORD INVISIBLE
user_entries[1].config(show='*')
user_entries[4].insert(0, '0')

# CREATING 'SETUP/RESET' CHECKBOX
chkValue = BooleanVar()  # If checkbox is true, run startup procedure. If false, run repair report procedure.
chkValue.set(False)
setup_checkbox = Checkbutton(root, text='Setup/Reset', command=update_repair_button_text, var=chkValue)

# CREATING BUTTONS
# Prompt user to find their desired directory where the report will be saved.
dir_button = Button(root, text='Search', command=directory_button, padx=80)
# dir_button.bind('<Button-1>', directory_button(user_info))
# load_button = Button(root, text='Load', command=load_info, padx=80)
# load_button.bind('<Button-1>', load_info(user_info, user_entries))
# save_button = Button(root, text='Save', command=save_info, padx=80)
# save_button.bind('<Button-1', save_info(user_info))
report_button = Button(root, text='Generate Report', command=checkbox_state, padx=55)
# SETTING BUTTONS POSITIONS
dir_button.grid(row=3, column=2)
report_button.grid(row=4, column=2)
# load_button.grid(row=6, column=0)
# save_button.grid(row=6, column=1)
setup_checkbox.grid(row=0, column=2)

# EVENT HANDLER
update_event = Event()

# MAKE INITIAL MESSAGE VISIBLE IN STATUS TEXTBOX
load_options()
load_xpaths()
load_info()
reset_status_box()
chkValue.set(False)
update_repair_button_text()


# MAIN LOOP
root.mainloop()
